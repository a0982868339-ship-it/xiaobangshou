from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_agreement():
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('婚内财产协议书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 甲方乙方信息
    p = doc.add_paragraph()
    p.add_run('甲方：').bold = True
    p.add_run(' 姓名，身份证号：____________________')
    
    p = doc.add_paragraph()
    p.add_run('乙方：').bold = True
    p.add_run(' 姓名，身份证号：____________________')
    
    doc.add_paragraph('甲乙双方于 ____ 年 ____ 月 ____ 日登记结婚。目前购房合同登记权利人为甲乙双方共有。为明确财产归属，双方经平等协商，达成如下协议：')
    
    # 第一条
    doc.add_heading('第一条 房产基本情况', level=1)
    doc.add_paragraph('1. 房产地址：________________________________________________')
    doc.add_paragraph('2. 购房合同编号/不动产权证号：________________________________')
    doc.add_paragraph('3. 该房产性质为安居型商品房。')
    
    # 第二条
    doc.add_heading('第二条 出资情况确认', level=1)
    doc.add_paragraph('1. 双方一致确认：该房产的全部购房款（包括但不限于首付款、余款、税费、契税、维修基金等）均由甲方父母以现金或银行转账方式全额支付。')
    doc.add_paragraph('2. 乙方确认：乙方未对该房产支付过任何购房款项，亦未参与后期的装修费用。')
    
    # 第三条
    doc.add_heading('第三条 房产归属约定', level=1)
    doc.add_paragraph('1. 双方一致约定：尽管该房产目前在购房合同/权属证书上登记为甲乙双方共有，但该房产的实际所有权 100% 归甲方个人所有，属于甲方的个人财产，不属于夫妻共同财产。')
    doc.add_paragraph('2. 乙方明确表示：自愿放弃该房产中其所占有的全部法律份额（包括目前及未来可能产生的增值部分）。')
    doc.add_paragraph('3. 甲方父母的出资行为仅视为对甲方个人的单方赠与，不属于对甲乙双方的赠与。')
    doc.add_paragraph('4. 若甲乙双方未来发生离婚，乙方承诺不就该房产主张任何权利，该房产不作为夫妻共同财产进行分割。')
    
    # 第四条
    doc.add_heading('第四条 权属登记变更', level=1)
    doc.add_paragraph('1. 待安居房政策允许办理更名或除名登记时，乙方应无条件配合甲方及甲方父母将房产权利人变更为甲方一人名下。')
    doc.add_paragraph('2. 办理变更登记所产生的税费由甲方承担。')
    
    # 第五条
    doc.add_heading('第五条 其他', level=1)
    doc.add_paragraph('1. 本协议自双方签字之日起生效。')
    doc.add_paragraph('2. 本协议一式三份，甲乙双方各执一份，另一份用于公证。')
    
    # 签署区
    doc.add_paragraph('\n\n')
    doc.add_paragraph('甲方（签字）：________________    日期：2026年__月__日')
    doc.add_paragraph('乙方（签字）：________________    日期：2026年__月__日')
    doc.add_paragraph('父母（见证签字）：________________')
    
    # 保存文档
    doc.save('婚内财产协议书_甲方版.docx')
    print("Word文档已成功生成：婚内财产协议书_甲方版.docx")

if __name__ == "__main__":
    create_agreement()
